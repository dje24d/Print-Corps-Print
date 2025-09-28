import os
import PyPDF2
from docx import Document
from PIL import Image
import io
try:
    import magic
    HAS_MAGIC = True
except ImportError:
    HAS_MAGIC = False
    import mimetypes

def count_pages(file_path, file_content=None):
    """
    Count pages in a document based on file type
    Returns page count or None if unable to determine
    """
    try:
        # Determine file type
        if HAS_MAGIC:
            if file_content:
                file_type = magic.from_buffer(file_content, mime=True)
            else:
                file_type = magic.from_file(file_path, mime=True)
        else:
            # Fallback to extension-based detection
            if file_content:
                filename = "temp_file"  # Placeholder
            else:
                filename = file_path
            _, ext = os.path.splitext(filename)
            ext = ext.lower()
            
            # Map extensions to MIME types
            extension_map = {
                '.pdf': 'application/pdf',
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.doc': 'application/msword',
                '.jpg': 'image/jpeg',
                '.jpeg': 'image/jpeg',
                '.png': 'image/png',
                '.tiff': 'image/tiff',
                '.tif': 'image/tiff',
                '.bmp': 'image/bmp',
                '.gif': 'image/gif'
            }
            file_type = extension_map.get(ext, 'unknown')
        
        # Handle different file types
        if 'pdf' in file_type:
            return count_pdf_pages(file_path, file_content)
        elif 'word' in file_type or 'officedocument' in file_type:
            return count_docx_pages(file_path, file_content)
        elif 'image' in file_type:
            return count_image_pages(file_path, file_content)
        else:
            return None
            
    except Exception as e:
        print(f"Error counting pages: {e}")
        return None

# Keep the count_pdf_pages, count_docx_pages, and count_image_pages functions the same

def count_pdf_pages(file_path, file_content=None):
    """Count pages in PDF file"""
    try:
        if file_content:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
        else:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
        return len(pdf_reader.pages)
    except Exception as e:
        print(f"Error reading PDF: {e}")
        return None

def count_docx_pages(file_path, file_content=None):
    """Estimate pages in DOCX file (approximate)"""
    try:
        if file_content:
            doc = Document(io.BytesIO(file_content))
        else:
            doc = Document(file_path)
        
        # This is an approximation as DOCX doesn't have fixed pages
        # Based on average words per page (250-300 words)
        total_words = 0
        for paragraph in doc.paragraphs:
            total_words += len(paragraph.text.split())
        
        # Estimate pages based on word count
        estimated_pages = max(1, round(total_words / 300))
        return estimated_pages
        
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        return None

def count_image_pages(file_path, file_content=None):
    """Count image pages (for multi-page TIFF)"""
    try:
        if file_content:
            image = Image.open(io.BytesIO(file_content))
        else:
            image = Image.open(file_path)
        
        # Check if it's a multi-page image (like TIFF)
        page_count = 0
        try:
            while True:
                page_count += 1
                image.seek(page_count)
        except EOFError:
            pass
        
        return max(1, page_count)  # At least 1 page
        
    except Exception as e:
        print(f"Error reading image: {e}")
        return None

def get_file_type(file_path, file_content=None):
    """Get file type using magic"""
    try:
        if file_content:
            return magic.from_buffer(file_content, mime=True)
        else:
            return magic.from_file(file_path, mime=True)
    except:
        return "unknown"